"""Narrative report — assemble EDA + analysis + recs into a Word document.

Public API:
- `build_report(out_path, ingest_result, eda_report, analyses, recs)`
    Writes a .docx report to `out_path`. Uses python-docx directly so the
    output works without the docx skill being loaded.

The report sections:
    1. Executive Summary (LLM narrative or stub)
    2. Dataset overview (rows/cols, schema, kind)
    3. Key findings (headline lines + correlations + outliers)
    4. Segments / Products (RFM, product matrix excerpts)
    5. Forecasts & risk (churn, stockout)
    6. Recommendations (sorted, with confidence + impact)
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

import pandas as pd

try:
    from docx import Document  # type: ignore
    from docx.shared import Pt, RGBColor, Inches
    HAS_DOCX = True
except Exception:  # pragma: no cover
    Document = None
    HAS_DOCX = False


def _add_table(doc, df: pd.DataFrame, max_rows: int = 12) -> None:
    df = df.head(max_rows)
    if df.empty:
        return
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, v in enumerate(row.values):
            cells[i].text = (f"{v:.2f}" if isinstance(v, float) else str(v))[:80]


def build_report(out_path: str | Path,
                 *,
                 ingest_result=None,
                 eda_report=None,
                 narrative: Optional[str] = None,
                 rfm_df: Optional[pd.DataFrame] = None,
                 matrix_df: Optional[pd.DataFrame] = None,
                 elasticity_df: Optional[pd.DataFrame] = None,
                 churn_df: Optional[pd.DataFrame] = None,
                 stockout_df: Optional[pd.DataFrame] = None,
                 basket_df: Optional[pd.DataFrame] = None,
                 recommendations: Optional[list] = None,
                 title: str = "Analyst Report") -> Path:
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed; run pip install python-docx.")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)

    # -- Executive summary ---------------------------------------------------
    doc.add_heading("Executive Summary", level=1)
    if narrative:
        doc.add_paragraph(narrative)
    elif eda_report and eda_report.headline:
        for line in eda_report.headline:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("(No summary provided.)")

    # -- Dataset overview ----------------------------------------------------
    if ingest_result is not None:
        doc.add_heading("Dataset overview", level=1)
        doc.add_paragraph(
            f"Rows: {ingest_result.n_rows:,}   "
            f"Columns: {ingest_result.n_cols}   "
            f"Detected kind: {ingest_result.kind.kind} "
            f"(confidence {ingest_result.kind.confidence:.0%})"
        )
        if ingest_result.issues:
            doc.add_paragraph("Data issues flagged:")
            for i in ingest_result.issues:
                doc.add_paragraph(i, style="List Bullet")

        schema_rows = []
        for c in ingest_result.schema.columns:
            schema_rows.append({
                "column": c.name,
                "type": c.inferred_type,
                "null %": f"{c.null_pct:.1%}",
                "unique %": f"{c.unique_pct:.1%}",
                "reason": c.reason,
            })
        if schema_rows:
            _add_table(doc, pd.DataFrame(schema_rows), max_rows=30)

    # -- Findings ------------------------------------------------------------
    if eda_report:
        doc.add_heading("Key findings", level=1)
        if eda_report.correlations:
            doc.add_paragraph("Top correlations:")
            for c in eda_report.correlations[:5]:
                doc.add_paragraph(f"{c.a} ↔ {c.b}: r = {c.pearson:+.2f}",
                                  style="List Bullet")
        if eda_report.outliers:
            doc.add_paragraph("Outliers detected:")
            for o in eda_report.outliers[:5]:
                doc.add_paragraph(f"{o.column} — {o.n_outliers} via {o.method} ({o.pct:.1%})",
                                  style="List Bullet")
        if eda_report.seasonality and eda_report.seasonality.notes:
            doc.add_paragraph("Seasonality:")
            for n in eda_report.seasonality.notes:
                doc.add_paragraph(n, style="List Bullet")

    # -- Customers + products ------------------------------------------------
    if rfm_df is not None and not rfm_df.empty:
        doc.add_heading("Customer segments (RFM)", level=1)
        seg_counts = rfm_df["segment"].value_counts().reset_index()
        seg_counts.columns = ["segment", "customers"]
        _add_table(doc, seg_counts)

    if matrix_df is not None and not matrix_df.empty:
        doc.add_heading("Product performance matrix", level=1)
        _add_table(doc, matrix_df[["product", "revenue", "share", "growth", "quadrant"]])

    if elasticity_df is not None and not elasticity_df.empty:
        doc.add_heading("Price elasticity", level=1)
        _add_table(doc, elasticity_df)

    # -- Risk ----------------------------------------------------------------
    if churn_df is not None and not churn_df.empty:
        doc.add_heading("Customer churn risk", level=1)
        risk_counts = churn_df["risk"].value_counts().reset_index()
        risk_counts.columns = ["risk", "customers"]
        _add_table(doc, risk_counts)

    if stockout_df is not None and not stockout_df.empty:
        doc.add_heading("Inventory risk", level=1)
        _add_table(doc, stockout_df)

    if basket_df is not None and not basket_df.empty:
        doc.add_heading("Cross-sell opportunities (basket)", level=1)
        _add_table(doc, basket_df)

    # -- Recommendations -----------------------------------------------------
    if recommendations:
        doc.add_heading("Recommended actions", level=1)
        for i, r in enumerate(recommendations, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {r.action}")
            run.bold = True
            doc.add_paragraph(f"Why: {r.evidence}", style="Intense Quote")
            doc.add_paragraph(f"Confidence: {r.confidence}   Impact: {r.impact_estimate}")
            doc.add_paragraph(f"Audience: {r.audience}   Category: {r.category}")
            doc.add_paragraph()  # spacer

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


__all__ = ["build_report", "HAS_DOCX"]
